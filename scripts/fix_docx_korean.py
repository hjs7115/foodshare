from docx import Document


DOCX_PATH = r"C:\Users\82102\OneDrive\Desktop\컴퓨터공학과\졸작\ver2\banding_merged\반띵_허준서_강신혁_소프트웨어설계서_수정본.docx"
OUT_PATH = r"C:\Users\82102\OneDrive\Desktop\컴퓨터공학과\졸작\ver2\banding_merged\반띵_허준서_강신혁_소프트웨어설계서_수정본_한글수정.docx"


def set_cell_text(cell, text):
    cell.text = text


def set_row(row, values):
    for cell, value in zip(row.cells, values):
        set_cell_text(cell, value)


def replace_paragraph_text(paragraph, new_text):
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


doc = Document(DOCX_PATH)

# 1. Scope
t = doc.tables[1]
set_cell_text(
    t.rows[1].cells[1],
    "회원가입/로그인, 게시글 기반 식재료 등록 및 관리, 유통기한 관리, 식재료 소분 거래, 공동구매 게시글 작성 및 참여, 댓글 기능, 반경 기반 거리 필터링, 거래 신청 및 상태 관리, 거래 내역 관리, 관심 게시글 관리, 사용자 신뢰도(평점) 기능",
)
set_cell_text(
    t.rows[2].cells[1],
    "실제 결제 시스템 연동, AI 기반 자동 식단 추천, 배달 서비스 연동, 모바일 네이티브 앱 개발(iOS/Android), 실시간 채팅 기능, 실시간 위치 추적 기능",
)

# 2. Functional requirements
t = doc.tables[2]
set_row(
    t.rows[2],
    [
        "FR-002",
        "게시글 기반 식재료 등록 및 관리",
        "사용자는 거래 게시글 작성 시 식재료명, 수량, 유통기한, 가격, 거래 장소, 이미지 등을 등록·수정·삭제할 수 있다.",
        "상",
        "게시글 중심 관리",
    ],
)
set_cell_text(
    t.rows[4].cells[2],
    "사용자는 공동구매 게시글을 작성하거나 기존 공동구매에 참여할 수 있으며, 모집 인원과 마감일을 관리할 수 있다.",
)
set_cell_text(
    t.rows[5].cells[2],
    "시스템은 게시글에 등록된 식재료의 유통기한을 관리하고, 임박한 식재료를 목록 및 알림 대상으로 표시한다.",
)
set_row(
    t.rows[11],
    [
        "FR-011",
        "관심 게시글 관리",
        "사용자는 관심 있는 소분 거래 및 공동구매 게시글을 저장하고 마이페이지에서 조회·해제할 수 있다.",
        "중",
        "마이페이지 연동",
    ],
)

# 2-1. Non-functional requirement note
t = doc.tables[3]
set_cell_text(t.rows[2].cells[4], "모바일 웹 기준")

# 3. Use case
t = doc.tables[5]
set_cell_text(
    t.rows[4].cells[1],
    "1. 회원은 게시글 등록 화면에 접근한다.\n2. 식재료 정보(이름, 수량, 가격, 유통기한, 거래 장소)를 입력한다.\n3. 거래 유형(나눔/판매/공동구매)을 선택한다.\n4. 게시글을 등록한다.\n5. 다른 회원이 게시글을 조회하고 거래 또는 공동구매 참여를 신청한다.\n6. 작성자가 신청을 수락하면 거래가 진행된다.",
)
set_cell_text(
    t.rows[6].cells[1],
    "거래 게시글과 참여 신청 정보가 저장되며, 거래 상태가 대기 상태로 관리된다.",
)

# 4. Modules
t = doc.tables[7]
set_row(t.rows[9], ["Favorite Module", "관심 게시글 등록, 해제, 조회 기능", "입력: 게시글 ID → 출력: 관심 목록"])
set_row(t.rows[10], ["Upload Module", "게시글 및 프로필 이미지 업로드 처리", "입력: 이미지 파일 → 출력: 이미지 URL"])

# 5. Database added rows
t = doc.tables[8]
db_rows = {
    44: ["users", "latitude", "DOUBLE", "Y", "", "사용자 위치 위도"],
    45: ["users", "longitude", "DOUBLE", "Y", "", "사용자 위치 경도"],
    46: ["posts", "target_count", "INT", "Y", "", "공동구매 목표 인원"],
    47: ["posts", "current_count", "INT", "Y", "", "공동구매 현재 참여 인원"],
    48: ["posts", "deadline", "DATETIME", "Y", "", "공동구매 모집 마감일"],
    49: ["posts", "updated_at", "DATETIME", "Y", "", "게시글 수정일시"],
    50: ["comments", "updated_at", "DATETIME", "Y", "", "댓글 수정일시"],
    51: ["favorites", "id", "BIGINT", "N", "PK", "관심 게시글 고유 ID (AUTO_INCREMENT)"],
    52: ["favorites", "user_id", "BIGINT", "N", "FK", "관심 등록 사용자 ID"],
    53: ["favorites", "post_id", "BIGINT", "N", "FK", "관심 등록 게시글 ID"],
    54: ["favorites", "created_at", "DATETIME", "N", "", "관심 등록일시"],
}
for row_idx, values in db_rows.items():
    set_row(t.rows[row_idx], values)

# 6. API descriptions
t = doc.tables[9]
api_rows = {
    9: ["GET", "/api/posts", "게시글 목록 조회", "Query: postType, keyword, sort, radius", "{ posts }"],
    10: ["POST", "/api/posts", "게시글 작성", "{ title, content, postType, quantity, price, expirationDate, tradeLocation, imageUrl }", "{ post }"],
    11: ["GET", "/api/posts/{postId}", "게시글 상세 조회", "-", "{ post }"],
    12: ["PUT", "/api/posts/{postId}", "게시글 수정", "{ title, content, quantity, price, status }", "{ post }"],
    13: ["DELETE", "/api/posts/{postId}", "게시글 삭제", "-", "{ success: true }"],
    14: ["GET", "/api/posts/{postId}/comments", "댓글 목록 조회", "-", "{ comments }"],
    15: ["POST", "/api/posts/{postId}/comments", "댓글 작성", "{ content }", "{ comment }"],
    16: ["PUT", "/api/comments/{commentId}", "댓글 수정", "{ content }", "{ comment }"],
    17: ["DELETE", "/api/comments/{commentId}", "댓글 삭제", "-", "{ success: true }"],
    18: ["POST", "/api/posts/{postId}/trade-requests", "거래 신청", "-", "{ tradeRequest }"],
    19: ["GET", "/api/posts/{postId}/trade-requests", "게시글별 거래 신청 조회", "-", "{ tradeRequests }"],
    20: ["POST", "/api/trade-requests/{tradeRequestId}/accept", "거래 신청 수락", "-", "{ tradeRequest }"],
    21: ["POST", "/api/trade-requests/{tradeRequestId}/reject", "거래 신청 거절", "-", "{ tradeRequest }"],
    22: ["GET", "/api/mypage", "마이페이지 프로필 조회", "-", "{ user }"],
    23: ["PUT", "/api/mypage", "프로필 수정", "{ nickname, profileImage, location }", "{ user }"],
    24: ["GET", "/api/mypage/posts", "내 게시글 조회", "-", "{ posts }"],
    25: ["GET", "/api/mypage/trade-requests", "내 거래 신청 내역 조회", "-", "{ tradeRequests }"],
    26: ["GET", "/api/mypage/received-trade-requests", "받은 거래 신청 조회", "-", "{ tradeRequests }"],
    27: ["POST", "/api/posts/{postId}/favorites", "관심 게시글 등록", "-", "{ favorite }"],
    28: ["DELETE", "/api/posts/{postId}/favorites", "관심 게시글 해제", "-", "{ success: true }"],
    29: ["GET", "/api/mypage/favorites", "관심 게시글 목록 조회", "-", "{ favorites }"],
    30: ["POST", "/api/uploads/images", "이미지 업로드", "multipart/form-data: image", "{ imageUrl }"],
    31: ["POST", "/api/users/{userId}/reviews", "후기 및 평점 등록", "{ tradeRequestId, rating, content }", "{ review }"],
    32: ["GET", "/api/users/{userId}/rating", "사용자 신뢰도 조회", "-", "{ rating, reviewCount }"],
}
for row_idx, values in api_rows.items():
    set_row(t.rows[row_idx], values)

# 7. External APIs / tech stack / UI rows
t = doc.tables[10]
set_cell_text(t.rows[3].cells[1], "거래 신청, 거래 수락, 유통기한 임박 등 푸시 알림 발송(향후 확장)")

t = doc.tables[11]
set_cell_text(t.rows[1].cells[1], "React + TypeScript + Vite")
set_cell_text(
    t.rows[1].cells[2],
    "컴포넌트 기반 UI 개발에 적합하며, Vite를 통해 빠른 개발 서버와 빌드 환경을 제공한다. 모바일 웹 화면을 우선 고려하여 앱과 유사한 사용자 경험을 제공한다.",
)
set_cell_text(t.rows[4].cells[1], "Docker, AWS EC2, ngrok(개발/시연)")
set_cell_text(
    t.rows[4].cells[2],
    "개발 및 배포 환경을 통일할 수 있으며, ngrok을 통해 개발 중인 백엔드 API를 프론트엔드와 임시 연동하여 시연할 수 있다.",
)

t = doc.tables[13]
set_row(
    t.rows[4],
    [
        "게시글 상세 화면",
        "",
        "사용자는 게시글의 상세 정보, 작성자 정보, 댓글, 관심 등록 여부를 확인할 수 있으며, 거래 신청 또는 공동구매 참여를 수행할 수 있다.",
    ],
)
set_row(
    t.rows[5],
    [
        "마이페이지 화면",
        "",
        "사용자는 프로필, 내 게시글, 관심 목록, 거래 내역, 알림 및 위치 설정, 평점 정보를 조회하고 관리할 수 있다.",
    ],
)

for p in doc.paragraphs:
    if "PWA" in p.text or "하단 네비게이션" in p.text:
        replace_paragraph_text(
            p,
            "본 시스템은 모바일 웹 환경을 우선 고려하여 설계되었으며, 하단 네비게이션을 통해 주요 기능 간 빠른 이동이 가능하도록 구성한다. 향후 필요 시 PWA 기능을 추가하여 홈 화면 설치 및 오프라인 캐시 기능으로 확장할 수 있다.",
        )
        break

t = doc.tables[14]
set_cell_text(t.rows[4].cells[2], "회원가입, 게시글, 댓글, 거래 신청 기능 구현")
set_cell_text(t.rows[5].cells[2], "거리 필터링, 관심 목록, 평점, 거래 내역 기능 구현")

# 8. Test cases
t = doc.tables[18]
test_rows = {
    2: ["TC-002", "게시글 기반 식재료 등록 및 관리", "1. 게시글 작성 화면 접속 → 2. 식재료명, 수량, 유통기한, 가격 입력 → 3. 저장 버튼 클릭", "등록한 식재료 거래 게시글이 목록에 표시되고 수정/삭제가 가능하다.", "FR-002 FR-005"],
    5: ["TC-005", "댓글 작성 및 조회", "1. 게시글 상세 화면 접속 → 2. 댓글 입력 → 3. 등록 후 댓글 목록 확인", "댓글이 정상 등록되고 해당 게시글 상세 화면에 표시된다.", "FR-008"],
    6: ["TC-006", "관심 게시글 등록/해제", "1. 게시글 상세 화면 접속 → 2. 관심 버튼 클릭 → 3. 마이페이지 관심 목록 확인 → 4. 관심 해제", "관심 목록에 게시글이 추가되고 해제 시 목록에서 제거된다.", "FR-011"],
    7: ["TC-007", "내 게시글 관리", "1. 마이페이지 접속 → 2. 내 게시글 목록 조회 → 3. 게시글 삭제", "사용자가 작성한 게시글만 조회되며 삭제 시 목록과 서버 데이터에 반영된다.", "FR-003 FR-009"],
    8: ["TC-008", "위치 반경 필터링", "1. 위치 설정 → 2. 반경 값 변경 → 3. 게시글 목록 조회", "설정 반경 내 게시글만 목록에 표시된다.", "FR-006"],
}
for row_idx, values in test_rows.items():
    set_row(t.rows[row_idx], values)

# 9. Risk management
t = doc.tables[19]
set_row(
    t.rows[4],
    [
        "4",
        "프론트엔드 localStorage 임시 저장 데이터와 백엔드 DB 데이터 불일치",
        "중",
        "상",
        "백엔드 API를 우선 저장소로 사용하고, localStorage는 캐시 또는 장애 시 임시 표시 용도로만 제한한다.",
    ],
)
set_row(
    t.rows[5],
    [
        "5",
        "외부 API 및 ngrok 주소 변경으로 인한 연동 오류",
        "중",
        "중",
        "API Base URL을 환경변수로 분리하고, 시연 전 서버 주소와 CORS 설정을 점검한다.",
    ],
)

for p in doc.paragraphs:
    if "2026-05-24부 1차 설계서 수정" in p.text:
        if "2026-05-26부 2차 설계서 수정" not in p.text:
            p.add_run("\n2026-05-26부 2차 설계서 수정. (현재 프론트엔드 구현 범위 반영, API 명세서 보강, DB 구조 및 테스트 케이스 보완)")
        break

doc.save(OUT_PATH)
print(OUT_PATH)
