from copy import deepcopy
from docx import Document


SRC_PATH = r"C:\Users\82102\OneDrive\Desktop\컴퓨터공학과\졸작\ver2\banding_merged\반띵_허준서_강신혁_소프트웨어설계서_수정본_한글수정.docx"
OUT_PATH = r"C:\Users\82102\OneDrive\Desktop\컴퓨터공학과\졸작\ver2\banding_merged\반띵_허준서_강신혁_소프트웨어설계서_수정본_PWA반영.docx"


def set_cell_text(cell, text):
    cell.text = text


def set_row(row, values):
    for cell, value in zip(row.cells, values):
        set_cell_text(cell, value)


def clone_append_row(table, values, template_index=-1):
    tr = deepcopy(table.rows[template_index]._tr)
    table._tbl.append(tr)
    row = table.rows[-1]
    set_row(row, values)
    return row


def replace_paragraph_text(paragraph, new_text):
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


doc = Document(SRC_PATH)

# 1. Scope: PWA should be in scope, not a future-only extension.
t = doc.tables[1]
set_cell_text(
    t.rows[1].cells[1],
    "회원가입/로그인, 게시글 기반 식재료 등록 및 관리, 유통기한 관리, 식재료 소분 거래, 공동구매 게시글 작성 및 참여, 댓글 기능, 반경 기반 거리 필터링, 거래 신청 및 상태 관리, 거래 내역 관리, 관심 게시글 관리, 사용자 신뢰도(평점) 기능, PWA 기반 설치형 웹앱 및 기본 오프라인 캐시",
)

# 2. Non-functional requirements: restore PWA as a main requirement.
t = doc.tables[3]
set_cell_text(t.rows[2].cells[4], "PWA 핵심 적용")
clone_append_row(
    t,
    [
        "NFR-008",
        "PWA",
        "서비스는 모바일 브라우저에서 설치 가능한 웹앱 형태를 지원하고, 기본 화면 및 정적 리소스에 대한 캐시를 제공해야 한다.",
        "앱 설치 가능, 기본 오프라인 화면 제공",
        "manifest/service worker 적용",
    ],
)

# 3. External API: add PWA browser APIs.
t = doc.tables[10]
clone_append_row(
    t,
    [
        "Web App Manifest / Service Worker",
        "홈 화면 설치, 정적 리소스 캐싱, 오프라인 fallback 화면 제공",
        "Browser API",
        "HTTPS 환경 필요",
    ],
)

# 4. Tech stack: PWA as a core frontend characteristic.
t = doc.tables[11]
set_cell_text(t.rows[1].cells[1], "React + TypeScript + Vite + PWA")
set_cell_text(
    t.rows[1].cells[2],
    "컴포넌트 기반 UI 개발과 빠른 빌드 환경을 제공하며, Web App Manifest와 Service Worker를 적용하여 모바일에서 설치 가능한 앱과 유사한 사용자 경험을 제공한다.",
)

# 5. PWA paragraph in UI/UX section.
for p in doc.paragraphs:
    if "본 시스템은 모바일 웹 환경을 우선 고려하여 설계되었으며" in p.text or "본 시스템은 PWA 기반으로 설계되어" in p.text:
        replace_paragraph_text(
            p,
            "본 시스템은 PWA 기반으로 설계되어 모바일 브라우저에서도 앱처럼 설치하여 사용할 수 있으며, 하단 네비게이션을 통해 주요 기능 간 빠른 이동이 가능하다. 또한 Web App Manifest와 Service Worker를 활용하여 기본 정적 리소스 캐싱과 오프라인 fallback 화면을 제공함으로써 모바일 환경에서 안정적인 사용자 경험을 제공한다.",
        )
        break

# 6. WBS: include PWA setup explicitly.
t = doc.tables[14]
set_cell_text(t.rows[4].cells[2], "회원가입, 게시글, 댓글, 거래 신청 기능 및 PWA 기본 설정 구현")
set_cell_text(t.rows[5].cells[2], "거리 필터링, 관심 목록, 평점, 거래 내역 기능 및 PWA 사용성 개선")

# 7. Milestones: include PWA in prototype and beta milestones.
t = doc.tables[15]
set_cell_text(t.rows[2].cells[4], "회원가입, 게시글 작성, 거래 신청 기능, PWA 설치 가능 환경 구현")
set_cell_text(t.rows[3].cells[4], "위치 기반 기능, 알림 기능, PWA 캐시 및 통합 테스트")

# 8. Test cases: add PWA behavior test.
t = doc.tables[18]
clone_append_row(
    t,
    [
        "TC-009",
        "PWA 설치 및 오프라인 기본 동작",
        "1. 모바일 브라우저에서 서비스 접속 → 2. 홈 화면 설치 가능 여부 확인 → 3. 네트워크 차단 후 기본 화면 접근",
        "서비스가 설치 가능한 웹앱으로 인식되며, 오프라인 상태에서도 기본 안내 화면 또는 캐시된 화면이 표시된다.",
        "NFR-008",
    ],
)

# 9. Risk management: add PWA/browser compatibility risk.
t = doc.tables[19]
clone_append_row(
    t,
    [
        "6",
        "브라우저별 PWA 설치 및 Service Worker 동작 차이",
        "중",
        "중",
        "Chrome과 Edge를 기준으로 우선 검증하고, HTTPS 환경에서 manifest와 service worker 등록 상태를 테스트한다.",
    ],
)

# 10. Revision history.
for p in doc.paragraphs:
    if "2026-05-26부 2차 설계서 수정" in p.text:
        if "2026-05-26부 3차 설계서 수정" not in p.text:
            p.add_run("\n2026-05-26부 3차 설계서 수정. (PWA를 프로젝트 핵심 특징으로 반영, 비기능 요구사항·기술 스택·테스트 케이스 보완)")
        break

doc.save(OUT_PATH)
print(OUT_PATH)
